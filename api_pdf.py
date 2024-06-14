from flask import Blueprint, request, jsonify
import fitz
import vertexai
from vertexai.preview.generative_models import GenerativeModel
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from google.oauth2 import service_account
from google.cloud import storage
from google.cloud import translate_v2 as translate
from datetime import datetime

load_dotenv()

ws_pdf = Blueprint('ws_pdf', __name__)

# Cargar credenciales desde el archivo JSON
CREDENTIALS_PATH = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
if not CREDENTIALS_PATH:
    raise ValueError("GOOGLE_APPLICATION_CREDENTIALS no está definido.")
credentials = service_account.Credentials.from_service_account_file(CREDENTIALS_PATH)

# Configuración de Vertex AI
PROJECT_ID = os.getenv("PROJECT_ID")
REGION = os.getenv("REGION")
if not PROJECT_ID or not REGION:
    raise ValueError("PROJECT_ID o REGION no están definidos.")
vertexai.init(project=PROJECT_ID, location=REGION, credentials=credentials)

# Configuración del modelo generativo
generative_multimodal_model = GenerativeModel('gemini-1.0-pro-002')

# Inicializar el cliente de Google Cloud Storage
storage_client = storage.Client(credentials=credentials)
BUCKET_NAME = os.getenv("GCS_BUCKET_NAME")
if not BUCKET_NAME:
    raise ValueError("GCS_BUCKET_NAME no está definido.")

bucket = storage_client.bucket(BUCKET_NAME)

# Inicializar el cliente de Google Translate
translate_client = translate.Client(credentials=credentials)

def translate_text(text, target_language='es'):
    result = translate_client.translate(text, target_language=target_language)
    return result['translatedText']

def format_text_to_word(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def format_markdown_text(text):
    # Formatear encabezados
    text = re.sub(r'^#+ ', lambda match: match.group().replace('#', ''), text, flags=re.MULTILINE)

    # Formatear cursiva
    text = re.sub(r'\*(.+?)\*', r'*\1*', text)

    # Formatear listas
    text = re.sub(r'^\* ', lambda match: '• ', text, flags=re.MULTILINE)

    return text

# Función para extraer texto entre dos páginas específicas
def extraer_texto_entre_paginas(doc, pagina_inicio, pagina_fin, titulo_inicio=None, titulo_fin=None):
    texto_total = ""
    for pagina_num in range(pagina_inicio, pagina_fin + 1):
        pagina_obj = doc.load_page(pagina_num - 1)
        texto = pagina_obj.get_text("text")
        if pagina_num == pagina_inicio and titulo_inicio:
            inicio = texto.find(titulo_inicio)
            if inicio != -1:
                texto = texto[inicio + len(titulo_inicio):]
        if pagina_num == pagina_fin and titulo_fin:
            fin = texto.find(titulo_fin)
            if fin != -1:
                texto = texto[:fin]
        texto_total += texto.strip() + " "
    return texto_total.strip()

@ws_pdf.route('/extraer', methods=['POST'])
def extraer_texto_pdf():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No hay archivo en la solicitud'}), 400

        file = request.files['file']
        indicaciones = request.form.get('indicaciones', '')

        if file.filename == '':
            return jsonify({'error': 'No se ha seleccionado ningún archivo'}), 400

        if file:
            filename = secure_filename(file.filename)
            if filename is None:
                raise ValueError("El nombre del archivo es None")

            blob = bucket.blob(f"documentos/{filename}")
            blob.upload_from_file(file)

            temp_pdf_path = f"/tmp/{filename}"
            blob.download_to_filename(temp_pdf_path)

            # Abrir el documento PDF
            try:
                doc = fitz.open(temp_pdf_path)
            except Exception as e:
                raise ValueError(f"Error al abrir el PDF: {str(e)}")

            subtitulos_texto = []
            has_toc = False

            try:
                # Obtener la tabla de contenido
                tabla_contenido = doc.get_toc()
                if tabla_contenido:
                    has_toc = True

                if not has_toc:
                    # Si no hay tabla de contenido, extraer texto de todas las páginas
                    for page_num in range(1, doc.page_count + 1):
                        texto = extraer_texto_entre_paginas(doc, page_num, page_num)
                        subtitulos_texto.append({'title': f'Página {page_num}', 'text': texto})
                else:
                    for i, entry in enumerate(tabla_contenido):
                        title = entry[1]
                        page_number = entry[2]
                        level = entry[0]

                        if level == 2:
                            if i + 1 < len(tabla_contenido):
                                next_entry = tabla_contenido[i + 1]
                                next_title = next_entry[1]
                                next_page_number = next_entry[2]
                            else:
                                next_title = None
                                next_page_number = doc.page_count

                            texto = extraer_texto_entre_paginas(doc, page_number, next_page_number, title, next_title)
                            subtitulos_texto.append({'title': title, 'text': texto})
            finally:
                doc.close()

            generated_content = []
            cantidad_peticiones = 0

            # Crear un documento de Word
            final_doc = Document()

            for subtitulo in subtitulos_texto:
                # Generar contenido
                response = generative_multimodal_model.generate_content([f"Genera la respuesta según las siguientes indicaciones: {indicaciones}\n\n{subtitulo['text']}"])
                generated_text = response.candidates[0].content.parts[0].text

                # Traducir el texto generado al idioma especificado (español)
                translated_text = translate_text(generated_text)

                # Formatear el texto traducido
                formatted_text = format_markdown_text(translated_text)

                # Agregar el título al documento de Word solo si hay tabla de contenido
                if has_toc:
                    paragraph = final_doc.add_heading(subtitulo['title'], level=1)
                    paragraph.style.font.size = Pt(14)
                    paragraph.style.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Agregar el texto traducido al documento de Word
                paragraph = final_doc.add_paragraph()
                for line in formatted_text.split('\n'):
                    format_text_to_word(paragraph, line)
                    paragraph.add_run('\n')

                generated_content.append(formatted_text)
                cantidad_peticiones += 1

            # Generar la marca de tiempo para el nombre del archivo
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{filename}_{timestamp}.docx"
            output_path = f"/tmp/{output_filename}"

            # Guardar el documento de Word temporalmente
            final_doc.save(output_path)

            # Subir el documento de Word a Google Cloud Storage
            blob = bucket.blob(f"documentos/{output_filename}")
            blob.upload_from_filename(output_path)

            # Obtener la URL pública del archivo
            public_url = f"https://storage.cloud.google.com/{BUCKET_NAME}/documentos/{output_filename}"

            return jsonify({
                'message': 'Documento generado con éxito',
                'output_url': public_url,
                'cantidad_peticiones': cantidad_peticiones
            }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500
